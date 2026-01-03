import sys
from docx import Document

def read_docx_structure(docx_path):
    try:
        document = Document(docx_path)
        print(f"Structure of {docx_path}:\n")
        
        count = 0
        for para in document.paragraphs:
            if not para.text.strip():
                continue
            
            style_name = "Normal"
            if para.style and hasattr(para.style, 'name'):
                style_name = para.style.name
            
            print(f"[{style_name}] {para.text.strip()[:100]}")
            
            count += 1
            if count > 20:
                break

            
            # If it's normal text, only print if it looks like a list item or start of section
            # to avoid flooding output
            # elif para.text.strip().startswith(('1.', '-', '*')):
            #    print(f"    {para.text.strip()[:50]}...")
            
        # Also try to read tables as they often contain requirements
        # print("\n--- TABLES ---\n")
        # for table in document.tables:
        #     for row in table.rows:
        #         row_text = [cell.text.strip() for cell in row.cells]
        #         print(f"| {' | '.join(row_text)} |")
                
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_docx.py <path>")
        sys.exit(1)
    read_docx_structure(sys.argv[1])
