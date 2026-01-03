import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_to_docx(md_path, docx_path):
    document = Document()
    
    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_block_content = []
    
    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()
        
        # Handle Code Blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End of code block
                p = document.add_paragraph()
                run = p.add_run('\n'.join(code_block_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.style = 'No Spacing' # Or just normal
                code_block_content = []
                in_code_block = False
            else:
                # Start of code block
                if in_table: # Close table if open
                    create_table(document, table_data)
                    table_data = []
                    in_table = False
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line.rstrip())
            continue

        # Handle Tables
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
            table_data.append(stripped)
            continue
        else:
            if in_table:
                create_table(document, table_data)
                table_data = []
                in_table = False

        if not stripped:
            continue

        # Handle Headings
        if stripped.startswith('#'):
            level = len(stripped.split(' ')[0])
            text = stripped[level:].strip()
            # docx only supports 1-9
            if level > 9: level = 9
            document.add_heading(text, level=level)
        
        # Handle Images
        elif stripped.startswith('![') and '](' in stripped:
            try:
                alt_text = stripped.split('![')[1].split(']')[0]
                img_path = stripped.split('](')[1].split(')')[0]
                # Resolve path relative to md file location
                full_img_path = os.path.join(os.path.dirname(md_path), img_path)
                
                # document.add_paragraph(f"[图片: {alt_text}]")
                if os.path.exists(full_img_path):
                    try:
                        document.add_picture(full_img_path, width=Inches(6))
                    except Exception as e:
                        print(f"Could not add image {full_img_path}: {e}")
                        document.add_paragraph(f"[Image: {alt_text} - File not found or format not supported]")
                else:
                    document.add_paragraph(f"[Image: {alt_text} - Placeholder]")
            except:
                 p = document.add_paragraph()
                 add_run_with_formatting(p, stripped)

        # Handle Horizontal Rule
        elif stripped.startswith('---'):
            document.add_page_break()

        # Normal Paragraph
        else:
            p = document.add_paragraph()
            add_run_with_formatting(p, stripped)

    if in_table:
        create_table(document, table_data)

    document.save(docx_path)
    print(f"Document saved to {docx_path}")

def create_table(document, table_lines):
    # Filter out separator lines like |---|---|
    # Regex to match lines that contain only | - : and spaces
    data_lines = [l for l in table_lines if not re.match(r'^\|[\s:-]+\|$', l.replace(' ', '')) and not re.match(r'^[\s:-]+\|$', l.replace(' ', ''))]
    
    # Specific check for markdown separator row: | --- | :--- |
    # We can check if the row only contains dashes, colons and pipes
    final_rows = []
    for l in data_lines:
        # A simple heuristic: if a cell contains letters/numbers, it's data. 
        # If it's just dashes and colons, it's a separator.
        if re.search(r'[a-zA-Z0-9\u4e00-\u9fa5]', l): # Check for alphanumeric or Chinese
            final_rows.append(l)
        else:
             # It might be a separator
             pass

    # Re-parsing: standard markdown table logic is: Header, Separator, Data...
    # My logic above was a bit loose. Let's try to assume the second line is separator if it looks like one.
    
    rows = []
    for l in table_lines:
        # Check if separator
        clean_l = l.strip()
        # Remove outer pipes
        if clean_l.startswith('|'): clean_l = clean_l[1:]
        if clean_l.endswith('|'): clean_l = clean_l[:-1]
        
        cells = [c.strip() for c in clean_l.split('|')]
        
        # Check if this row is a separator (---)
        is_sep = True
        for c in cells:
            if not re.match(r'^[\s:-]+$', c):
                is_sep = False
                break
        
        if is_sep:
            continue
            
        rows.append(cells)
    
    if not rows:
        return

    # Create table
    num_cols = max(len(r) for r in rows)
    if num_cols == 0: return

    table = document.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                # Bold headers
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

def add_run_with_formatting(paragraph, text):
    # Simple bold parser for **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python script.py input.md output.docx")
        sys.exit(1)
    
    parse_markdown_to_docx(sys.argv[1], sys.argv[2])
